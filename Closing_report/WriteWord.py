# Import docx NOT python-docx
import docx
import json
from docx.shared import Pt
from docx.shared import Inches
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH

with open('details.json', 'r+') as f:
        json_data = json.load(f)

Dept_name = json_data['dept_name']
Acad_year = json_data['Acad_yar']
Prog_name = json_data['Prog_name']
Sem = json_data['Sem']
Course_code_and_name = json_data['Course_name'] + ' ' + json_data['Course_details']['Course_CODE']
Course_cordinator = json_data['Course_details']['Course_Cordinator']
Course_Outcomes = json_data['COS']
doc = docx.Document()

doc.add_heading(f'Name of the department: {Dept_name}', 0)

y = doc.add_heading(f'AY: {Acad_year}', 3)
y.alignment = 1

y = doc.add_heading('Course Closing Report: ', 2)
y.alignment = 1

doc.add_heading(f'Programe Name: {Prog_name}', 3)

doc.add_heading(f'Semester: {Sem}', 3)

doc.add_heading(f'Course Name & Code: {Course_code_and_name}\n', 3)

para = doc.add_paragraph().add_run('Course Outcomes:')
para.font.size = Pt(12)
para.font.underline = True
para.font.bold = True

doc.add_heading('At the completion of the course, students will be able to\n', 3)

def table_maker(doc, COS_all):
    table = doc.add_table(rows=1, cols=3)

    hdr_cells = table.rows[0].cells
    run = hdr_cells[0].paragraphs[0].add_run('S.No')
    run.bold = True
    # run.underline = True
    # table.autofit = False 

    run = hdr_cells[1].paragraphs[0].add_run('COURSE OUTCOMES')
    run.bold = True
    # run.underline = True

    run = hdr_cells[2].paragraphs[0].add_run('COGNITIVE LEVELS')
    run.bold = True

    for i in COS_all:
        SNO = i
        Course_outcome = COS_all[i][0]
        Cognitive_level = COS_all[i][1]

        hdr_cells = table.add_row().cells

        hdr_cells[0].text = str(SNO)
        hdr_cells[1].text = str(Course_outcome)
        hdr_cells[2].text = str(Cognitive_level)

    table.allow_autofit = False
    for cell in table.columns[0].cells:
        cell.width = Inches(1)
    for cell in table.columns[1].cells:
        cell.width = Inches(3.5)
    for cell in table.columns[2].cells:
        cell.width = Inches(2)
    
        
COS_all = json_data['COS'][0]
rows = len(COS_all)
table_maker(doc, COS_all)

para = doc.add_paragraph().add_run('\nCO-PO and CO-PSO Mapping:')
para.font.size = Pt(12)
para.font.underline = True
para.font.bold = True

para = doc.add_paragraph().add_run('\nCO Attainments in 2019-20:')
para.font.size = Pt(12)
para.font.underline = True
para.font.bold = True

para = doc.add_paragraph().add_run('\nPO-PSO Attainments in 2019-20:')
para.font.size = Pt(12)
para.font.underline = True
para.font.bold = True

para = doc.add_paragraph().add_run('\nSummary of Result Analysis:')
para.font.size = Pt(12)
para.font.underline = True
para.font.bold = True

y = doc.add_heading('Signature: \t\t\t\t\t\t\t\t\t Signature: ', 2)
y.alignment = 0

y = doc.add_heading('Module Coordinator: \t\t\t\t\t\t\t Course Coordinator: ', 3)
y.alignment = 0

para = doc.add_paragraph().add_run(f'{Course_cordinator}')
para.font.size = Pt(12)
para.font.underline = True
para.font.bold = True

# y = doc.add_heading(f'{Course_cordinator}', 2)
# y.alignment = 0

# y = doc.add_heading('Module Coordinator: \t\t\t\t\t\t\t Course Coordinator: ', 3)
# y.alignment = 0

# Adding paragraph with normal font size
# doc.add_heading('Normal Font Size Paragraph:', 3)
# doc.add_paragraph(
# 	'GeeksforGeeks is a Computer Science portal for geeks.')

# y = doc.add_heading(f'Course Closing Report', 4)
# y.add_run()
# title_style = y.style
# title_style.font.name = "Times New Roman"

# Now save the document to a location
doc.save('1. Closing report.docx')
os.system('libreoffice "1. Closing report.docx"')
