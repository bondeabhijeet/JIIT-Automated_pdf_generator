import tkinter as tk
from tkinter import ttk
from tkinter import *
import docx
from docx.shared import Pt
from matplotlib.colors import cnames
import entry as ET
import table_maker as TM
import json
import COs as COs
# INIT
doc = docx.Document()

root = tk.Tk()
root.title('Closing Report')
root.geometry('1200x800')

def json_maker(data):
    json_object = json.dumps(data, indent = 4)

    with open("details.json", "w") as outfile:
        outfile.write(json_object)

# For getting values from the entries and doing furthur operations
def get_value():
    Dept_name = dept_name.get()
    Academic_Year = Aca_year.get()
    Programme_name = Prog_name.get()
    Semester = Sem.get()
    Course_name = C_name.get()
    No_of_COS = no_of_cos.get()

    print(Dept_name, Academic_Year, Programme_name, Semester, Course_name, No_of_COS)

    doc.add_heading(f'Name of the department: {Dept_name}', 0)
    y = doc.add_heading(f'AY: {Academic_Year} ({radioValue})', 3)
    y.alignment = 1

    y = doc.add_heading(f'Course Closing Report', 2)
    y.alignment = 1

    doc.add_heading(f'Programe Name: {Programme_name}', 4)
    doc.add_heading(f'Semester: {Semester}', 4)
    doc.add_heading(f'Course Name: {Course_name}\n\n', 4)

    para = doc.add_paragraph().add_run('Course Outcomes:')
    para.font.size = Pt(13)
    para.font.underline = True
    para.font.bold = True

    doc.add_heading('At the completion of the course, students will be able to\n', 4)

    TM.table_maker(doc)
    
    data = {'dept_name':f'{Dept_name}','Acad_yar':f'{Academic_Year} ({radioValue})', 'Prog_name':f'{Programme_name}', 'Sem':f'{Semester}', 'Course_name':f'{Course_name}', 'COS':[]}
    json_maker(data)

    COs.COS(int(No_of_COS))

    doc.save('main.docx')
    # root.destroy()

def radio_get():
    global radioValue
    radioValue = str(v.get())
    print(radioValue)

# Creating a radio button 
def radio(root, radioOptions):
    global v
    v = StringVar(root, "1")
    
    values = radioOptions

    for (text, value) in values.items():
        Radiobutton(root, text = text, variable = v,
                    value = value, indicator = 0,
                    background = "light blue", command=radio_get).pack(fill = X, ipady = 5)

def label_handler(text):
    label = ttk.Label(text=text)
    label.pack(fill=tk.X, padx=5, pady=5)

label_handler("Enter academic year")
Aca_year = ET.entry(root)
radio(root, {"OddSem" : "Odd Semester", "EvenSem" : "Even Semester"})

label_handler("Programme Name")
Prog_name = ET.entry(root)

label_handler("Semester")
Sem = ET.entry(root)

label_handler("Course Name")
C_name = ET.entry(root)

ttk.Label(root, text='Department name').pack(fill=tk.X ,padx=5, pady=5)
dept_name = ttk.Entry(root)
dept_name.pack(fill=tk.X, padx=5, pady=5)

ttk.Label(root, text="Number of Course Outcome's").pack(fill=tk.X ,padx=5, pady=5)
no_of_cos = ttk.Entry(root)
no_of_cos.pack(fill=tk.X, padx=5, pady=5)

butt1 = ttk.Button(text="SAVE CHANGES AND FILL CO's", command=get_value)
butt1.pack(fill=tk.X, padx=5, pady=5)

root.mainloop()