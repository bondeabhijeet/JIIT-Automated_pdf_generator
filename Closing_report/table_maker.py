# # Import docx NOT python-docx
# import docx

# # Create an instance of a word document
# doc = docx.Document()

# # Add a Title to the document
# doc.add_heading('GeeksForGeeks', 0)

# # Table data in a form of list
# data = (
# 	(1, 'Geek 1'),
# 	(2, 'Geek 2'),
# 	(3, 'Geek 3')
# )

# # Creating a table object
# table = doc.add_table(rows=4, cols=2)

# table.cell(0,0).width = 1097280
# table.cell(1,0).width = 1097280
# table.cell(2,0).width = 1097280


# table.cell(0,1).width = 4846320
# table.cell(1,1).width = 4846320
# table.cell(2,1).width = 4846320

# # Adding heading in the 1st row of the table
# row = table.rows[0].cells
# row[0].text = 'Id'
# row[1].text = 'Name'

# # Adding data from the list to the table
# for id, name in data:

# 	# Adding a row and then adding data in it.
# 	row = table.add_row().cells
# 	row[0].text = str(id)
# 	row[1].text = name

# # Adding style to a table
# table.style = 'Colorful List'

# # Now save the document to a location
# doc.save('gfg.docx')

import docx
from docx.shared import Inches

# doc = docx.Document()
# doc.add_heading('Name: ', level=1)
def table_maker(doc):
    table = doc.add_table(rows=4, cols=3)

    hdr_cells = table.rows[0].cells
    run = hdr_cells[0].paragraphs[0].add_run('S.No')
    run.bold = True
    # run.underline = True
    # table.autofit = False 

    run = hdr_cells[1].paragraphs[0].add_run('COURSE OUTCOMES')
    run.bold = True
    # run.underline = True

    run = hdr_cells[2].paragraphs[0].add_run('COGNITIVE LEVELS')
    run.bold = True

    table.allow_autofit = False
    for cell in table.columns[0].cells:
        cell.width = Inches(1)
    for cell in table.columns[1].cells:
        cell.width = Inches(3.5)
    for cell in table.columns[2].cells:
        cell.width = Inches(2)
# table.cell(0,0).text = 'Time Zone'
# table.cell(1,0).text = 'Link'
# table.cell(1,1).text = 'https://www.google.com/'
# table.cell(2,0).text = 'Website'
# table.cell(3,0).text = 'Facebook'

# doc.save('test.docx')